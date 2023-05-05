# Script de automatizat documentul pentru BD - Laborator
# Creaza in mod automat un document word cu numarul laboratorului ca titlu centrat si paragrafe de tip "Exercitiul " + index pentru fiecare exercitiu in parte.
# Paragrafele by default sunt scrise cu "Helvetica" , Bold de 12 delimitate de 2 newline-uri pentru a oferi spatiu rezolvarilor.
# Partile sunt cu font de 20, in stil Heading si cu underline
# Facut de Mat(2.718)i din lenea de a reface documentul la fiecare laborator de Baze de Date.

import docx
from docx.shared import Pt
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def createDocument(nrLaborator, nrExercitii, nrParti):
    document = docx.Document()
    titlu = document.add_heading("Rezolvare Laborator "+str(nrLaborator), 0)
    titlu.alignment = 1
    if (nrParti == 1):
        for counter in range(1, nrExercitii+1):
            paragraph = document.add_paragraph("Exercitiul "+str(counter))
            paragraph.style.font.size = Pt(12)
            paragraph.style.font.name = 'Helvetica'
            paragraph.style.font.bold = True
            run = paragraph.add_run()
            run.add_break()
            run.add_break()
            paragraph.style.font.bold = False
        document.save("D:\Programare\AutomateBD\Laborator_" +
                      str(nrLaborator)+".docx")
    elif (nrParti > 1):
        for partiCounter in range(1, nrParti+1):
            print("Cum se numeste aceasta parte?")
            parte = input()
            print("Cate exercitii are partea > "+parte+" < ?")
            nrExercitiiParte = int(input())
            partiHeading = document.add_heading(str(parte), 1)
            partiHeading.style.font.size = Pt(20)
            partiHeading.style.font.underline = True
            for counter in range(1, nrExercitiiParte+1):
                paragraph = document.add_paragraph("Exercitiul "+str(counter))
                paragraph.style.font.size = Pt(12)
                paragraph.style.font.name = 'Helvetica'
                paragraph.style.font.bold = True
                run = paragraph.add_run()
                run.add_break()
                run.add_break()
            document.save("D:\Programare\AutomateBD\Laborator_" +
                          str(nrLaborator)+".docx")


def callCreateDocument():
    print("""
  _                                    __   __                  
 /_)     _)_  _   _ _   _  _)_   _     )_)  ) )     )   _  ( _  
/ /  (_( (_  (_) ) ) ) (_( (_   )_)   /__) /_/     (__ (_(  )_) 
                               (_               __              
    """)
    print("Bine ai venit la AutomateBD! \nAcest script simplu in python iti creaza in mod automat un document Word cu formatul cerut pentru laboratoare, deci nu mai trebuie sa stai de fiecare data sa pierzi timpul.")
    print("Totul este foarte simplu : Introduci numarul laboratorului. Dupa numarul de parti ale acestuia.\nIntrodu 1 pentru un singur set de exercitii. In cazul in care sunt mai multe, pune mai multe idk.\nIn final, introdu numarul de exercitii pentru fiecare parte.")
    print("""__  __  __  __  __  __                    
    """)
    print("Ce numar are laboratorul?")
    nrLaborator = int(input())
    print("Cate parti are laboratorul?")
    nrParti = int(input())
    if (nrParti == 1):
        print("Cate exercitii are laboratorul?")
        nrExercitii = int(input())
        createDocument(nrLaborator, nrExercitii, nrParti)
        print("""__  __  __  __  __  __                    
        """)
        print('Document creat! •ᴗ•')
    elif (nrParti > 1):
        createDocument(nrLaborator, 1, nrParti)
        print("""__  __  __  __  __  __                    
        """)
        print('Document creat! •ᴗ•')
    else:
        print("Numarul de parti trebuie sa fie cel putin 1!")
        return
    print("""
                                       /  __      ___       __  \    
 _   _   _|  _   |_       |\/|  _  |_ (    _)       /  /|  (__)  ) . 
||| (_| (_| (-   |_) \/   |  | (_| |_  \  /__ .    /    |  (__) /  | 
                     /                                               
        """)


callCreateDocument()
